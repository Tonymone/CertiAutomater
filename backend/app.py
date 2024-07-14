from flask import Flask, request, send_file, jsonify
import pandas as pd
import numpy as np
import os
import cv2
from docx import Document
from docx.shared import Inches
from docx2pdf import convert
import tempfile
from flask_cors import CORS
import time
import shutil

app = Flask(__name__)
CORS(app)  # Enable CORS for all routes
UPLOAD_FOLDER = 'uploads'
GEN_FOLDER = 'gens'
TEMPLATE_PATH = 'certificate-template.jpg'

# Set the Downloads folder path (change the path as needed)
DOWNLOADS_FOLDER = os.path.join(os.path.expanduser("~"), 'Downloads')

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['GEN_FOLDER'] = GEN_FOLDER

status = {"message": ""}  # Declare the status variable globally

if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)

if not os.path.exists(GEN_FOLDER):
    os.makedirs(GEN_FOLDER)

def delete_files_in_folder(folder_path):
    for filename in os.listdir(folder_path):
        file_path = os.path.join(folder_path, filename)
        try:
            if os.path.isfile(file_path) or os.path.islink(file_path):
                os.unlink(file_path)
            elif os.path.isdir(file_path):
                shutil.rmtree(file_path)
        except Exception as e:
            print(f'Failed to delete {file_path}. Reason: {e}')

@app.route('/status', methods=['GET'])
def get_status():
    global status
    return jsonify(status)

@app.route('/delete-files', methods=['POST'])
def delete_files():
    try:
        delete_files_in_folder(app.config['GEN_FOLDER'])
        delete_files_in_folder(app.config['UPLOAD_FOLDER'])
        return "Files deleted successfully", 200
    except Exception as e:
        print(f"Error deleting files: {e}")
        return "Error deleting files", 500

@app.route('/generate-certificates', methods=['POST'])
def generate_certificates():
    # delete_files_in_folder('gens')
    global status

    if 'ms6File' not in request.files or 'bmsFile' not in request.files:
        return "No file part", 400

    ms6_file = request.files['ms6File']
    bms_file = request.files['bmsFile']

    ms6_file_path = os.path.join(app.config['UPLOAD_FOLDER'], 'MS6.xlsx')
    bms_file_path = os.path.join(app.config['UPLOAD_FOLDER'], 'BMS.xlsx')

    ms6_file.save(ms6_file_path)
    bms_file.save(bms_file_path)

    df1 = pd.read_excel(ms6_file_path)
    df2 = pd.read_excel(bms_file_path)

    df2['FREM'].fillna('null', inplace=True)
    df2['RES'].fillna('null', inplace=True)

    dataT = df2[(df2['RSLT'] == 'P') & (df2['FREM'] == 'null') & (df2['RES'] == 'null')]
    dataT['Gender'] = 'null'
    dataT = pd.merge(dataT, df1[['COLL_NO', 'COLL_NAME']], on='COLL_NO', how='left')
    dataT = dataT.sort_values(by='COLL_NO', ascending=True)
    dataT['pno'] = dataT.groupby('COLL_NO').cumcount() + 1
    dataT['COLL_NO'] = dataT['COLL_NO'].apply(lambda x: str(x).zfill(4))
    dataT['pno'] = dataT['pno'].apply(lambda x: str(x).zfill(4))

    # Generate certificates
    try:
        status['message'] = "Generating certificates..."
        for index, row in dataT.iterrows():
            certificate_template_image = cv2.imread(TEMPLATE_PATH)
            
            name = str(row['NAME']).strip() if pd.notnull(row['NAME']) else 'N/A'
            coll_name = str(row['COLL_NAME']).strip() if pd.notnull(row['COLL_NAME']) else 'N/A'
            coll_no = str(row['COLL_NO']) + ' : '
            pno = str(row['pno'])
            
            cv2.putText(certificate_template_image, name, (815, 1500), cv2.FONT_HERSHEY_SIMPLEX, 2, (0, 0, 250), 5, cv2.LINE_AA)
            cv2.putText(certificate_template_image, coll_name, (815, 1700), cv2.FONT_HERSHEY_SIMPLEX, 2, (0, 0, 250), 5, cv2.LINE_AA)
            cv2.putText(certificate_template_image, coll_no, (2575, 490), cv2.FONT_HERSHEY_SIMPLEX, 2, (0, 0, 250), 5, cv2.LINE_AA)
            cv2.putText(certificate_template_image, pno, (2850, 490), cv2.FONT_HERSHEY_SIMPLEX, 2, (0, 0, 250), 5, cv2.LINE_AA)
            
            output_path = os.path.join(app.config['GEN_FOLDER'], f"{name}.jpg")
            cv2.imwrite(output_path, certificate_template_image)
            print(f"Generated certificate: {output_path}")

        # Create Word document
        status['message'] = "Generating Word document..."
        output_word_path = os.path.join(app.config['GEN_FOLDER'], "certificates.docx")
        create_word_document(dataT, output_word_path)

        # Convert to PDF
        status['message'] = "Generating PDF file..."
        output_pdf_path = os.path.join(app.config['GEN_FOLDER'], "certificates.pdf")
        convert(output_word_path, output_pdf_path)

        status['message'] = "Completed"
        response = send_file(output_pdf_path, as_attachment=True, download_name="certificates.pdf")
    
    except Exception as e:
        print(f"Error generating certificates: {e}")
        status['message'] = "Error generating certificates"
        return "Error generating certificates", 500
    
    # finally:
        # Delete files after 3 seconds
        # time.sleep(3)
        # delete_files_in_folder('gens')
        # delete_files_in_folder('uploads')
    
    return response

def create_word_document(data, output_path):
    document = Document()
    certificate_folder = app.config['GEN_FOLDER']

    sections = document.sections
    for section in sections:
        section.top_margin = Inches(0)
        section.bottom_margin = Inches(0)
        section.left_margin = Inches(0.82)
        section.right_margin = Inches(0)

    for i in range(0, len(data), 2):
        row1 = data.iloc[i]
        certificate1 = os.path.join(certificate_folder, f"{str(row1['NAME']).strip()}.jpg")

        if os.path.exists(certificate1):
            paragraph = document.add_paragraph()
            run = paragraph.add_run()
            run.add_picture(certificate1, width=Inches(6.87))

        if i + 1 < len(data):
            row2 = data.iloc[i + 1]
            certificate2 = os.path.join(certificate_folder, f"{str(row2['NAME']).strip()}.jpg")

            if os.path.exists(certificate2):
                paragraph = document.add_paragraph()
                run = paragraph.add_run()
                run.add_picture(certificate2, width=Inches(6.87))
                document.add_page_break()
            else:
                print(f"File not found: {certificate2}")
        else:
            document.add_page_break()

    document.save(output_path)

if __name__ == '__main__':
    app.run(debug=True)
